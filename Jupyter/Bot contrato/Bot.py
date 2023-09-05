import telegram
from telegram.ext import Updater, CommandHandler, MessageHandler, ConversationHandler, filters


from docx import Document
import docx2pdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
# Abre el archivo Word
documento = Document("Contrato Plantilla.docx")

# Estados de la conversación
GENDER, NAME, CEDULA, DATE, START_TIME, END_TIME, SIGN_DATE = range(7)

# Diccionario para la conversión de cédula
cedula_conversion = {
    "-": "Guion",
    "0": "Cero",
    "1": "Uno",
    "2": "Dos",
    "3": "Tres",
    "4": "Cuatro",
    "5": "Cinco",
    "6": "Seis",
    "7": "Siete",
    "8": "Ocho",
    "9": "Nueve"
}

# Función para convertir una cédula en letras
def convertir_cedula_a_letras(cedula):
    letras = []
    for caracter in cedula:
        if caracter in cedula_conversion:
            letras.append(cedula_conversion[caracter])
        else:
            letras.append(caracter)
    return " ".join(letras)

# Función para iniciar la conversación
def start(update, context):
    reply_text = "Hola, soy tu bot de contrato. Por favor, elige tu género: (1) Hombre, (2) Mujer"
    update.message.reply_text(reply_text)
    return GENDER

# Función para manejar la elección de género
def choose_gender(update, context):
    gender = int(update.message.text)
    context.user_data['gender'] = gender
    if gender == 1:
        context.user_data['saludo'] = "El Modelo"
        context.user_data['seno'] = "señor"
        context.user_data['gen'] = "hombre"
        context.user_data['pana'] = "panameño"
    elif gender == 2:
        context.user_data['saludo'] = "La Modelo"
        context.user_data['seno'] = "señora"
        context.user_data['gen'] = "mujer"
        context.user_data['pana'] = "panameña"
    else:
        update.message.reply_text("Opción no válida. Por favor, elige (1) Hombre o (2) Mujer.")
        return GENDER

    reply_text = f"Perfecto, {context.user_data['saludo']} {context.user_data['seno']}. Ahora, por favor ingresa tu nombre:"
    update.message.reply_text(reply_text)
    return NAME

# Función para manejar el nombre
def get_name(update, context):
    context.user_data['name'] = update.message.text
    reply_text = f"Gracias, {context.user_data['name']}. Ahora ingresa tu cédula con guiones:"
    update.message.reply_text(reply_text)
    return CEDULA

# Función para manejar la cédula
def get_cedula(update, context):
    cedula_input = update.message.text
    cedula_numerica = "".join(filter(str.isdigit, cedula_input))
    cedula_en_letras = convertir_cedula_a_letras(cedula_numerica) + f" ({cedula_numerica})"
    context.user_data['cedula'] = cedula_en_letras

    reply_text = "Ingresa la fecha en que se realizará la sesión (por ejemplo, sábado 9 de septiembre de 2023):"
    update.message.reply_text(reply_text)
    return DATE

# Función para manejar la fecha
def get_date(update, context):
    context.user_data['date'] = update.message.text
    reply_text = "Ingresa la hora de inicio (por ejemplo, 10:00 a.m.):"
    update.message.reply_text(reply_text)
    return START_TIME

# Función para manejar la hora de inicio
def get_start_time(update, context):
    context.user_data['start_time'] = update.message.text
    reply_text = "Ingresa la hora de finalización (por ejemplo, 12:00 p.m.):"
    update.message.reply_text(reply_text)
    return END_TIME

# Función para manejar la hora de finalización
def get_end_time(update, context):
    context.user_data['end_time'] = update.message.text
    reply_text = "Ingresa el día de la firma del contrato (por ejemplo, martes 5 de septiembre de 2023):"
    update.message.reply_text(reply_text)
    return SIGN_DATE

# Función para manejar la fecha de firma
def get_sign_date(update, context):
    context.user_data['sign_date'] = update.message.text
    reply_text = "¡Gracias! Procesando la información y generando el contrato..."

    # Diccionario con las variables a reemplazar
    variables_a_reemplazar = {
        "NOMBREM": context.user_data['name'],
        "CEDULAM": context.user_data['cedula'],
        "DIAS": context.user_data['date'],
        "HORAI": context.user_data['start_time'],
        "HORAF": context.user_data['end_time'],
        "DIAF": context.user_data['sign_date'],
        "NUMEROCEDULA": "".join(filter(str.isdigit, context.user_data['cedula'])),
        "señora": context.user_data['seno'],
        "mujer": context.user_data['gen'],
        "panameña": context.user_data['pana']
    }

     # Recorre el contenido del documento
    for paragraph in documento.paragraphs:
        for variable, valor in variables_a_reemplazar.items():
            if variable in paragraph.text:
                nuevo_texto = paragraph.text.replace(variable, valor)
                paragraph.clear()  # Borra el párrafo actual
                paragraph.add_run(nuevo_texto)  # Agrega el nuevo texto al párrafo

    # Guarda el documento con los cambios
    docx_filename = f"Contrato para {context.user_data['name']}.docx"
    documento.save(docx_filename)

    # Convierte el documento Word a PDF
    pdf_filename = docx_filename.replace(".docx", ".pdf")
    docx2pdf.convert(docx_filename, pdf_filename)

    # Envía el PDF al usuario
    context.bot.send_document(chat_id=update.message.chat_id, document=open(pdf_filename, 'rb'))

    update.message.reply_text("¡Listo! El contrato ha sido generado en PDF y enviado. ¡Que tengas un buen día!")

    # Cierra el archivo
    documento.close()

    return ConversationHandler.END
# Función principal para iniciar el bot
def main():
    # Reemplaza 'TU_TOKEN' con el token de tu bot de Telegram
    TOKEN = '6481717517:AAFcvHiH6RPlIb5hJW4sL6h0iQ0gHANlTf4'

    # Crea una instancia de Updater
    updater = Updater(token=TOKEN, use_context=True)

    # Crea un manejador de conversación
    conversation_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            GENDER: [MessageHandler(Filters.text & ~Filters.command, choose_gender)],
            NAME: [MessageHandler(Filters.text & ~Filters.command, get_name)],
            CEDULA: [MessageHandler(Filters.text & ~Filters.command, get_cedula)],
            DATE: [MessageHandler(Filters.text & ~Filters.command, get_date)],
            START_TIME: [MessageHandler(Filters.text & ~Filters.command, get_start_time)],
            END_TIME: [MessageHandler(Filters.text & ~Filters.command, get_end_time)],
            SIGN_DATE: [MessageHandler(Filters.text & ~Filters.command, get_sign_date)]
        },
        fallbacks=[]
    )

    dispatcher = updater.dispatcher
    dispatcher.add_handler(conversation_handler)

    # Inicia el bot
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
