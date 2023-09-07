#Variables
#NOMBREM=Nombre de la modelo.
#CEDULAM=Cedula de la modelo en letra y numero.
#DIAS=Fecha en que se realizara la sesion.
#HORAI=Hora de inicio.
#HORAF=Hora de finalizacion.
#DIAF=Dia de la firma del contrato.
#NUMEROCEDULA=Cedula en numero de la modelo.

from docx import Document

# Abre el archivo Word
documento = Document("Contrato Plantilla.docx")  # Reemplaza "mi_documento.docx" con la ruta de tu propio archivo

print("Ingrese el Genero: (1)Hombre , (2) Mujer")
genero=input
if genero == 1:
    impresion="El Modelo"
    seno="señor"
    gen="hombre"
    pana="panameño"
elif genero== 2:
    impresion="La Modelo"
    seno="señora"
    gen="mujer"
    pana="panameña"
    
print("Ingrese el Nombre de {}:".format(impresion))
NOMBREM=input
print("Ingrese la Cedula de {} con guiones:".format(impresion))
CEDULAMN=input
print("Ingrese la Fecha en que se realizara la sesion:")
DIAS=input
print("Ingrese la Hora de inicio:")
HORAI=input
print("Ingrese la Hora de finalizacion:")
HORAF=input
print("Ingrese el Dia de la firma del contrato :")
DIAF=input

cedula_conversion={
    "-":"Guion",
    "0":"Cero",
    "1":"Uno",
    "2":"Dos",
    "3":"Tres",
    "4":"Cuatro",
    "5":"Cinco",
    "6":"Seis",
    "7":"Siete",
    "8":"Ocho",
    "9":"Nueve" 
}
# Función para convertir una cédula en letras
def convertir_cedula_a_letras(cedula):
    letras = []
    for caracter in cedula:
        if caracter in cedula_conversion:
            letras.append(cedula_conversion[caracter])
        else:
            letras.append(caracter)
    return " ".join(letras)

CEDULAM=convertir_cedula_a_letras(CEDULAMN)+" ({})".format(CEDULAMN)
print (CEDULAMN)

# Diccionario con las variables a buscar y sus valores correspondientes
variables_a_reemplazar = {
    "NOMBREM": NOMBREM,
    "CEDULAM": CEDULAM,
    "DIAS": DIAS+" del 2023",
    "HORAI": HORAI,
    "HORAF": HORAF,
    "DIAF": DIAF,
    "NUMEROCEDULA": CEDULAMN,
    "señora":seno,
    "mujer":gen,
    "panameña":pana    
}

# Recorre el contenido del documento
for paragraph in documento.paragraphs:
    for variable, valor in variables_a_reemplazar.items():
        if variable in paragraph.text:
            nuevo_texto = paragraph.text.replace(variable, valor)
            paragraph.clear()  # Borra el párrafo actual
            paragraph.add_run(nuevo_texto)  # Agrega el nuevo texto al párrafo

# Guarda el documento con los cambios
documento.save("Contrato para {}.docx".format(NOMBREM))  # Cambia el nombre del archivo según sea necesario

# Cierra el archivo
documento.close()

print("Las variables han sido reemplazadas en el documento.")
